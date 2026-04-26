from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)
AMBER  = RGBColor(0xE8, 0x91, 0x3A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
MUTED  = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x2D, 0x6A, 0x4F)
RED    = RGBColor(0xDC, 0x26, 0x26)
LGRAY  = RGBColor(0xF3, 0xF4, 0xF6)

# ── Helper: set cell shading ──────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                tag.set(qn(f'w:{attr}'), val)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

# ── Helper: heading ───────────────────────────────────────────────────────────
def add_heading(text, level=1, color=None):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = color or NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # bottom border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '1B2A4A')
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = color or NAVY
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = color or AMBER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text, bold=False, italic=False, color=None, size=11, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size       = Pt(size)
    run.bold            = bold
    run.italic          = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(text, bold_prefix=None, color=None):
    p   = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        if color:
            r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    return p

def add_divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'E8913A')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(2)

# ── Helper: 2-col info table ──────────────────────────────────────────────────
def add_info_table(rows_data, col_widths=(2.2, 4.3)):
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    for label, value in rows_data:
        row    = t.add_row()
        c0, c1 = row.cells
        c0.width = Inches(col_widths[0])
        c1.width = Inches(col_widths[1])
        shade_cell(c0, 'EBF0F8')
        shade_cell(c1, 'FFFFFF')
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size  = Pt(10)
        r0.font.color.rgb = NAVY
        p0.paragraph_format.space_after  = Pt(2)
        p0.paragraph_format.space_before = Pt(2)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(10)
        p1.paragraph_format.space_after  = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Helper: full-width header table (navy bg) ─────────────────────────────────
def add_section_banner(text, sub=None):
    t   = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shade_cell(cell, '1B2A4A')
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = WHITE
    p.paragraph_format.space_after  = Pt(0 if not sub else 2)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if sub:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(sub)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xB0,0xBE,0xD4)
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.space_before = Pt(0)
    set_row_height(t.rows[0], 1.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Helper: feature card table ────────────────────────────────────────────────
def add_feature_table(headers, rows_data, col_widths=None):
    ncols = len(headers)
    t     = doc.add_table(rows=1, cols=ncols)
    t.style = 'Table Grid'
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, '1B2A4A')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if col_widths:
            cell.width = Inches(col_widths[i])
    # data rows
    for ri, row_data in enumerate(rows_data):
        row = t.add_row()
        bg  = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            shade_cell(cell, bg)
            p  = cell.paragraphs[0]
            if isinstance(val, tuple):
                r = p.add_run(val[0])
                r.bold = val[1]
                r.font.size = Pt(10)
                if len(val) > 2:
                    r.font.color.rgb = val[2]
            else:
                r = p.add_run(val)
                r.font.size = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if col_widths:
                cell.width = Inches(col_widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_callout(text, color_hex='FFF3CD', border_hex='E8913A'):
    t    = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shade_cell(cell, color_hex)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.italic = True
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("MOMENTO")
r.bold = True
r.font.size = Pt(48)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("The Generative City-Wallet")
r2.font.size = Pt(18)
r2.font.color.rgb = AMBER
r2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Complete Project Overview for Antigravity Development Team")
r3.font.size = Pt(13)
r3.font.color.rgb = MUTED

doc.add_paragraph().paragraph_format.space_after = Pt(10)

add_info_table([
    ("Hackathon",      "Hack-Nation × World Bank Youth Summit · Global AI Hackathon 2026"),
    ("Challenge",      "Challenge 01 — Generative City-Wallet by DSV-Gruppe"),
    ("Product Name",   "MOMENTO"),
    ("Build Target",   "Progressive Web App (PWA) — Next.js 14 or React + Vite"),
    ("Deployment",     "Vercel — Single command deploy, live URL for judges"),
    ("AI Model",       "claude-sonnet-4-6 via Anthropic API"),
    ("Document Role",  "Complete technical, UX, and business specification for development"),
], col_widths=(2.0, 4.5))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS MOMENTO
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 1 — WHAT IS MOMENTO", "The concept, the problem, and the core innovation")

add_heading("One-Line Definition", 2)
add_callout(
    '"MOMENTO is a real-time commercial intelligence layer that detects when a merchant needs '
    'a customer, finds a receptive person nearby, and generates a hyper-personalized, '
    'time-limited offer — before the user even knew they wanted it."'
)

add_heading("The Problem MOMENTO Solves", 2)
add_body(
    "Imagine walking through Stuttgart on a cold Tuesday morning. A café 80 meters away just "
    "brewed fresh coffee. Their morning rush never came — they are sitting at 40% of their normal "
    "transaction volume. You are walking slowly, phone in hand, clearly not commuting. You have "
    "15 minutes free. Nobody connects these two facts."
)
add_body(
    "Instead, your phone shows: '10% off at a restaurant you have never visited, valid for 30 days.' "
    "This is not a technology problem. All the data exists. The café's quiet period is measurable. "
    "Your intent is inferable. The weather is known. What is missing is the connective tissue — "
    "the layer that reads all these signals simultaneously and generates one specific, well-timed offer."
)

add_heading("The Three Failures MOMENTO Fixes", 2)
add_feature_table(
    ["Failure", "What Exists Today", "What MOMENTO Does"],
    [
        ("Static Offers", "10% off valid 30 days — same coupon regardless of context", "Offer generated only at the exact moment it is needed"),
        ("No Algorithmic Power", "Corner café has no personalization tools — competes with chalkboard vs Amazon", "Amazon-level predictive intelligence for every local merchant"),
        ("Context Blindness", "User is 80m from a quiet café and nothing connects them", "Triple Clock detects the alignment and fires automatically"),
    ],
    col_widths=[1.5, 2.8, 2.2]
)

add_heading("The Core Innovation — Triple Clock Theory", 2)
add_body(
    "An offer fires ONLY when three clocks align simultaneously. When any one is absent — MOMENTO waits. "
    "This is the competitive moat. Anyone can know it is cold. Anyone can know a café is nearby. "
    "Only MOMENTO knows all three facts simultaneously at the exact right moment."
)

add_feature_table(
    ["Clock", "What It Detects", "Data Source", "Trigger Condition"],
    [
        ("Clock 1\nMerchant Quiet", "The café needs a customer RIGHT NOW", "Simulated Payone transaction velocity", "Current transactions 40%+ below historical baseline"),
        ("Clock 2\nUser Intent", "This person is receptive RIGHT NOW", "On-device behavioral model (demo: React state + localStorage)", "Slow movement + browsing pattern + available time window"),
        ("Clock 3\nCity Ambient", "City conditions support this RIGHT NOW", "OpenWeatherMap API + public event feeds", "Cold temperature + high rain probability + relevant time slot"),
    ],
    col_widths=[1.3, 2.0, 2.0, 2.2]
)

add_heading("Brand Philosophy", 2)
add_body("MOMENTO — from Latin momentum (the right moment) and Spanish momento (instant).")
add_bullet("Every offer is a moment, not a marketing message", "Brand rule:")
add_bullet('"This moment has passed." — not "Offer expired"', "Expiry says:")
add_bullet('"Another moment is coming." — not "Offer rejected"', "Dismissal says:")
add_bullet("QR code unfolds from the card — no copy needed", "Acceptance:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — THREE STAKEHOLDERS
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 2 — THE THREE STAKEHOLDERS", "Who uses MOMENTO, what they experience, and what they get")

add_heading("Stakeholder 1 — Mia (The User)", 2)
add_body("28 years old. Marketing professional. Walking through Stuttgart on her lunch break.")

add_feature_table(
    ["What Mia Does", "What Mia Gets"],
    [
        ("Nothing differently — just walks through the city", "ONE relevant offer at the exact moment she can use it"),
        ("Does not open any app to browse", "No notification spam — maximum 3 per day"),
        ("Does not search for offers", "Discovers places she would never have searched for"),
        ("Does not clip coupons", "Feels like the city is working for her"),
        ("Optionally shares behavioral context signals", "Earns data dividend credit (€0.12) per redemption"),
    ],
    col_widths=[3.2, 3.3]
)
add_callout("Data Dividend: Mia's behavioral context has commercial value to merchants. Instead of invisible data extraction, MOMENTO gives her a direct, transparent share of that value. This is the opposite of how most apps work.")

add_heading("Stakeholder 2 — Hans (The Corner Café Owner)", 2)
add_body("Runs Café Müller in Stuttgart. No marketing budget. No tech team. No time for campaigns.")

add_body("Hans sets three things ONCE and never touches it again:", bold=True)
add_bullet("What he offers: Hot drinks and pastries")
add_bullet("Maximum discount he is comfortable with: 20%")
add_bullet("His goal: Fill Tuesday quiet hours between 10am and 1pm")

add_body("What happens automatically after that:", bold=True)
add_bullet("MOMENTO detects when his café goes quiet via Payone transaction signals")
add_bullet("Generates a contextual offer matching his three rules")
add_bullet("Finds receptive users in the nearby area")
add_bullet("Surfaces the offer on their phones")
add_bullet("Processes redemption through his existing point of sale terminal")
add_bullet("Shows him performance data on his dashboard in real time")

add_body("What Hans pays:", bold=True)
add_callout("Zero upfront. Zero monthly subscription. Zero per offer generated. Zero per offer dismissed. Hans pays a 3-5% success fee ONLY when a customer walks in and redeems. If MOMENTO sends zero customers, Hans pays zero. The system is self-funding from the revenue it generates.")

add_heading("Stakeholder 3 — DSV Gruppe (The Platform Owner)", 2)
add_body("DSV Gruppe owns Payone (payment processor) and is part of the German Savings Banks Financial Group (Sparkassen). They sit at the intersection of everything MOMENTO needs.")

add_feature_table(
    ["Asset DSV Owns", "What It Gives MOMENTO"],
    [
        ("Payone — payment processor for tens of thousands of German merchants", "Real-time transaction velocity data — the quiet period signal no competitor has"),
        ("Sparkassen — 370 regional banks, 50M customers, 200+ year trust relationships", "Pre-built merchant trust, zero cold-call merchant onboarding, existing user base"),
        ("Existing merchant POS hardware", "Redemption works through existing terminals — zero new hardware needed"),
        ("Existing Sparkasse banking app (50M installs)", "MOMENTO launches as a feature update — no cold-start download problem"),
        ("GDPR-compliant financial infrastructure", "Privacy architecture built on trusted banking-grade systems"),
    ],
    col_widths=[3.2, 3.3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 3 — SYSTEM ARCHITECTURE", "How every component connects, what each module does")

add_heading("What We Are Building: A PWA", 2)
add_body(
    "MOMENTO is a Progressive Web App. This means it is a website that behaves like a native app. "
    "The user visits momento.app in their mobile browser, taps Add to Home Screen, and it appears "
    "on their home screen like a real app. No App Store. No approval process. No install friction."
)
add_feature_table(
    ["PWA Capability", "Technical Implementation", "Why It Matters"],
    [
        ("Home screen install", "Web App Manifest with icons and display: standalone", "Judges can demo on any device via URL"),
        ("Push notifications", "Web Push API with VAPID keys via Service Worker", "Core product feature — offers arrive without app open"),
        ("Offline support", "Workbox service worker with cache strategies", "Preferences saved even without connection"),
        ("Sensor access", "Browser DeviceMotion and Geolocation APIs", "Intent detection without native app"),
        ("Instant deployment", "Vercel — one command, live URL immediately", "No App Store review delays"),
    ],
    col_widths=[1.8, 2.5, 2.2]
)

add_heading("Full Technology Stack", 2)
add_feature_table(
    ["Layer", "Technology", "Purpose"],
    [
        ("Frontend", "React + TypeScript + Vite", "Component-based UI, fast build"),
        ("PWA", "vite-plugin-pwa + Workbox", "Service worker, offline support, push notifications"),
        ("Styling", "Tailwind CSS", "Rapid UI with design tokens"),
        ("Animations", "Framer Motion", "Offer card accept, dismiss, expire transitions"),
        ("Backend", "Node.js + Express", "API server for all endpoints"),
        ("Real-time", "Socket.io", "Live merchant dashboard updates"),
        ("AI", "Claude API — claude-sonnet-4-6", "Offer parameter generation — returns pure JSON"),
        ("Payone", "Simulated via Node.js setInterval", "Realistic transaction velocity feed for demo"),
        ("Weather", "OpenWeatherMap API (free tier)", "Real weather data for Stuttgart"),
        ("Location", "Browser Geolocation API", "Four permission tiers, graceful degradation"),
        ("QR Codes", "qrcode npm package", "Dynamic one-time-use redemption tokens"),
        ("Token Security", "HMAC-SHA256 via Node.js crypto", "Unforgeable, single-use QR tokens"),
        ("Database", "In-memory Maps and arrays", "No database setup needed for hackathon demo"),
        ("Deployment", "Vercel", "One command, automatic HTTPS, global CDN"),
    ],
    col_widths=[1.3, 2.2, 3.0]
)

add_heading("Module 1 — Context Sensing Layer", 2)
add_body(
    "This module continuously evaluates whether the Triple Clock is aligned. It aggregates signals "
    "from three sources and produces a ContextState object that either triggers or waits."
)

add_heading("Payone Transaction Simulation", 3)
add_body(
    "In production, Payone sends real-time payment events. For the hackathon, we simulate this using "
    "a Node.js setInterval that generates realistic transaction patterns per merchant. Each merchant "
    "has a baseline — the average number of transactions in a given 15-minute window on a given day "
    "of the week. When the current velocity drops below 60% of that baseline, a QuietPeriodSignal "
    "is emitted and Clock 1 turns green."
)
add_bullet("Café Müller baseline on Tuesday 11am: 12 transactions per 15 minutes")
add_bullet("If current count is 4 or fewer — quiet period signal fires")
add_bullet("Signal emitted via Server-Sent Events to the frontend Triple Clock panel")
add_bullet("Merchant dashboard shows live velocity chart vs baseline")

add_heading("User Intent Detection (On-Device, Demo: Simulated)", 3)
add_body(
    "IMPORTANT FOR DEVELOPERS: In the hackathon build, user intent is simulated via demo controls. "
    "There is no real TensorFlow Lite model running. The demo controls panel (visible at ?demo=true) "
    "contains a toggle labeled 'Simulate: User is Browsing' that forces the intent clock green. "
    "In production this would be a TF Lite model using real accelerometer data."
)
add_bullet("Demo implementation: React state variable toggled via demo controls panel")
add_bullet("Three possible states: receptive-browsing, focused-commuting, stationary-unavailable")
add_bullet("Personal preference learning: localStorage weight updates after each offer interaction")
add_bullet("Production vision: TensorFlow Lite on-device model — described to judges as the real path")

add_heading("City Ambient Signals", 3)
add_body("Weather from OpenWeatherMap API. Polled every 10 minutes. Cached in memory.")
add_bullet("Trigger condition: feels-like temperature below 12 degrees Celsius")
add_bullet("Trigger condition: rain probability above 60 percent")
add_bullet("Event data from Eventbrite or simulated local events for demo")
add_bullet("Time-of-day patterns: lunch window, post-work window, weekend drift")

add_heading("Configuration Without Code Changes", 3)
add_body(
    "All signal thresholds live in context.config.json. Changing city or data source requires "
    "only a config edit, not a code change. This directly addresses the brief's requirement for "
    "a configurable system. Different merchant types, cities, or threshold values slot in as "
    "configuration."
)

add_heading("Module 2 — Generative Offer Engine", 2)
add_body(
    "When all three clocks align, the generative pipeline fires. This module calls the Claude API "
    "to produce structured offer parameters — not final copy. The copy is generated on the device."
)

add_heading("What the Claude API Receives", 3)
add_body("The API call contains ONLY these fields — no personal data, no raw location:")
add_bullet("Weather context: feels-like temperature, condition string, rain probability")
add_bullet("Merchant context: category, anonymized quiet score, average transaction value")
add_bullet("User intent bundle: state enum, mobility mode enum, estimated free minutes, city district string")
add_bullet("Merchant rules: maximum discount percentage, offer types, merchant name")
add_callout("CRITICAL: The Claude API never receives GPS coordinates, personal identifiers, device IDs, preference history, or any raw sensor data. Zero personal data reaches the cloud LLM. This is enforced in code, not policy.")

add_heading("What the Claude API Returns — Structured Parameters Only", 3)
add_body("Claude returns pure JSON with these fields. The model string is claude-sonnet-4-6.")
add_feature_table(
    ["Parameter", "Example Value", "Used For"],
    [
        ("discount", "15", "Discount percentage to apply — minimum effective, not maximum"),
        ("featuredProduct", "Cappuccino + Croissant", "Product to highlight in the offer"),
        ("headlineTone", "warm_comfort", "Tone selector for on-device copy generation"),
        ("urgencyLevel", "gentle", "Controls layout weight and expiry window"),
        ("expiryMinutes", "14", "How long the offer remains valid"),
        ("visualMood", "warm_amber", "Controls hero image selection and color palette"),
        ("colorPrimary", "#C4783A", "Primary color for the generated widget"),
        ("emotionalHook", "cold_weather_warmth", "Template key for Service Worker headline"),
        ("ctaText", "Warm me up", "Call to action button text — maximum 3 words"),
        ("brandedEnding.dismiss", "Another moment is coming.", "Exact dismiss copy — never change this"),
        ("brandedEnding.expire", "This moment has passed.", "Exact expiry copy — never change this"),
    ],
    col_widths=[1.8, 2.0, 2.7]
)

add_heading("On-Device Copy Generation", 3)
add_body(
    "The parameters are sent to the user's device. The on-device model (localStorage-based for the "
    "hackathon) generates the final headline copy using the parameters and the user's local preference "
    "weights. This copy is NEVER generated by the server. This is the privacy enforcement mechanism."
)
add_bullet("Template map: emotional hook identifier maps to a headline template")
add_bullet("cold_weather_warmth → 'Cold outside? Your {product} is waiting.'")
add_bullet("rain_incoming → 'Rain on the way. Warm up at {merchant}.'")
add_bullet("cozy_invitation → '{freeMinutes} minutes to spare? There is a quiet corner nearby.'")
add_bullet("discovery → 'There is something worth stopping for — {distance}m away.'")
add_body("After each interaction, the on-device model updates localStorage weights:")
add_feature_table(
    ["Interaction", "Weight Change", "Stored In"],
    [
        ("Offer redeemed quickly (under 30 seconds)", "+0.15 to category weight and tone weight", "localStorage only"),
        ("Offer redeemed slowly", "+0.05 to category weight and tone weight", "localStorage only"),
        ("Offer dismissed by user", "-0.10 to category weight and tone weight", "localStorage only"),
        ("Offer expired unseen", "-0.03 to tone weight only", "localStorage only"),
    ],
    col_widths=[2.5, 2.5, 1.5]
)
add_callout("These weights NEVER sync to the server. No API call is made on any of these events. The personal preference model lives entirely on the user's device and is deleted when the user clears their data.")

add_heading("Module 3 — Seamless Checkout and Redemption", 2)

add_heading("The Redemption Flow", 3)
add_feature_table(
    ["Step", "What Happens", "Technical Detail"],
    [
        ("1. User taps Claim", "QR code generated server-side", "POST /api/offer/:id/accept"),
        ("2. Token created", "HMAC-SHA256 signed one-time UUID", "Includes: merchantId, discount, expiresAt, sessionHash"),
        ("3. QR displayed", "Full-screen QR code rendered on phone", "qrcode npm package, countdown timer visible"),
        ("4. Merchant scans", "POS web app scans the QR", "POST /api/redemption/validate"),
        ("5. Token validated", "Server checks: valid HMAC, not expired, not reused, merchant matches", "Invalid = error message, valid = proceed"),
        ("6. Discount applied", "Payone simulates discounted transaction", "Webhook fires to MOMENTO backend"),
        ("7. Dashboard updates", "Merchant sees +1 redeemed in real time", "Socket.io emit to merchant dashboard"),
        ("8. User confirmed", "Redemption confirmation notification sent", "Type 3 push notification within 5 seconds"),
    ],
    col_widths=[1.5, 2.5, 2.5]
)

add_heading("QR Token Security", 3)
add_bullet("HMAC-SHA256 signature prevents forgery — cannot produce valid token without the secret")
add_bullet("Single-use enforcement — used token IDs stored in memory, second scan returns error")
add_bullet("Server-side expiry — expiry timestamp in payload, server validates against current time")
add_bullet("Screenshot sharing does not work — system detects already-used token immediately")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — UX AND DESIGN
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 4 — UX AND DESIGN", "Every screen, every interaction, and the four mandatory UX questions")

add_heading("Design System", 2)
add_feature_table(
    ["Token Name", "Hex Value", "Used For"],
    [
        ("Brand Navy", "#1B2A4A", "Primary brand color, headers, navigation, all backgrounds"),
        ("Brand Amber", "#E8913A", "CTA buttons, accent elements, active clock indicators"),
        ("Warm Amber", "#C4783A", "Offer card primary color in warm mood state"),
        ("Urgency Red", "#E85A3A", "Countdown timer when under 2 minutes remaining"),
        ("Success Green", "#2D6A4F", "Redeemed state confirmation"),
        ("Dismiss Grey", "#9B9B9B", "Dismiss and expire states"),
        ("Background", "#F8F5F0", "App off-white background — warm, not clinical"),
        ("Card White", "#FFFFFF", "Offer card background"),
        ("Text Primary", "#1B2A4A", "All main body text"),
        ("Text Muted", "#6B7280", "Merchant name, distance, secondary information"),
    ],
    col_widths=[1.5, 1.5, 3.5]
)

add_body("Typography rules:", bold=True)
add_bullet("Headline: font-bold text-2xl leading-tight — conversational, not ad-like")
add_bullet("Subtext: font-normal text-base — warm and supporting")
add_bullet("Merchant info: font-normal text-sm text-muted — de-emphasized, emotion sells not logo")
add_bullet("Countdown: font-mono font-bold — precise, creates urgency feel")
add_bullet("CTA button: font-semibold text-lg — inviting, not aggressive")

add_heading("The Four Mandatory UX Questions", 2)
add_body("The hackathon brief explicitly requires answers to these four questions in the demo. Every design decision must address them.")

add_heading("Q1: Where Does the Interaction Happen?", 3)
add_body("Primary: Lock-screen widget that appears on natural phone lift during detected browsing window. Secondary: In-app card when user opens the PWA directly. NOT a push notification that interrupts mid-task.")
add_body("In PWA context: Web Push notification fires. User taps it. PWA opens to the offer card.")

add_heading("Q2: How Does the Offer Address the User?", 3)
add_body("Always emotional-situational. Never factual-informative. Always on-device generated. Never server-written.")
add_feature_table(
    ["Wrong (Factual)", "Right (Emotional-Situational)"],
    [
        ("15% off at Café Müller, 300m away", "Cold outside? Your cappuccino is waiting."),
        ("Tuesday Special: 20% off hot drinks", "Twelve minutes to spare? There is a warm seat nearby."),
        ("Café Müller: Quiet hour discount", "Rain on the way. A quiet corner and a fresh brew — 3 mins away."),
    ],
    col_widths=[3.0, 3.5]
)

add_heading("Q3: What Happens in the First 3 Seconds?", 3)
add_body("The user must understand the entire offer without reading any body copy. Three-zone layout enforced:")
add_feature_table(
    ["Zone", "Content", "Percentage of Card"],
    [
        ("Top zone", "Hero image only — no text overlay — warm, mood-matched", "Top 40%"),
        ("Middle zone", "ONE bold headline maximum 8 words — emotional hook", "Middle 35%"),
        ("Bottom zone", "Merchant name · distance · countdown timer · ONE CTA button", "Bottom 25%"),
    ],
    col_widths=[1.5, 3.5, 1.5]
)
add_callout("Zero-reading test: Cover the text. Does the user still understand the offer from the image and single headline alone? If not, the design failed.")

add_heading("Q4: How Does the Offer End?", 3)
add_feature_table(
    ["Ending", "Animation", "Copy Shown"],
    [
        ("Accept", "QR code unfolds with spring scale reveal from within the card", "No copy — QR takes over the card entirely"),
        ("Dismiss (swipe left)", "Card slides off left edge of screen with easeIn", "Another moment is coming. — fades in center, disappears after 2s"),
        ("Expire (timer zero)", "Card desaturates to grey and fades to 30% opacity over 1.5s", "This moment has passed. — remains visible until user navigates"),
    ],
    col_widths=[1.5, 2.8, 2.2]
)
add_body("Framer Motion variant names for developers:", bold=True)
add_bullet("Accept: qrVariants — scale spring from 0 to 1, delay 0.1s")
add_bullet("Dismiss: dismissVariants — x: -110%, opacity 0, duration 0.3s, easeIn")
add_bullet("Expire: expireVariants — opacity 0.3, filter grayscale 100%, duration 1.5s")
add_bullet("Card entrance: cardVariants — y: 100 to 0, spring damping 20")

add_heading("All Application Screens", 2)

add_heading("Onboarding Flow — 5 Steps", 3)
add_body("Seeds the local preference model. ALL data stored in localStorage ONLY. Zero API calls during onboarding.")
add_feature_table(
    ["Screen", "Question", "Options"],
    [
        ("Step 1", "I prefer...", "Coffee / Tea / Both"),
        ("Step 2", "I shop for...", "Experiences / Products / Food and Drink / All"),
        ("Step 3", "I respond to...", "Discovery: hidden gem nearby | Value: 15% off | Comfort: warm drink on a cold day"),
        ("Step 4", "My free time slots...", "Mornings / Lunch / Evenings / Weekends — multi-select"),
        ("Step 5", "My city area...", "Free text input — this is the ONLY location used for Tier 3 users"),
        ("Step 6", "Location permission", "Show 4-tier choice screen BEFORE OS permission dialog"),
    ],
    col_widths=[0.8, 2.0, 3.7]
)

add_heading("Location Permission — Four-Tier Screen", 3)
add_body("Show this as an explicit user choice, not just an OS permission dialog:")
add_feature_table(
    ["Tier", "What User Grants", "What System Uses", "Experience"],
    [
        ("Tier 1", "Precise GPS", "Real-time location (500m zone only transmitted)", "Distance-precise offers with exact walking time"),
        ("Tier 2", "Approximate location", "District-level area only", "Neighborhood-level offers without precise distance"),
        ("Tier 3", "Nothing — manual city input", "Only the text string user typed in Step 5", "City-wide merchant broadcasts — no distance shown"),
        ("Tier 4", "Nothing at all", "Nothing", "Browse offers manually inside app — no location features"),
    ],
    col_widths=[0.8, 1.8, 2.0, 1.9]
)
add_callout("ABSOLUTE RULE: In Tier 3 and Tier 4, there is NO IP address lookup, NO WiFi SSID reading, NO network-based location inference of any kind. If the user said no, the system hears no. This is enforced in the locationService.ts code, not just in documentation.")

add_heading("User Home Screen", 3)
add_bullet("Idle state: shows city ambient data, current time, gentle brand message")
add_bullet("Active offer state: full offer card takes the screen")
add_bullet("Bottom navigation: Home, My Moments (history), My Data (privacy controls)")

add_heading("Offer Card Screen", 3)
add_bullet("Full-screen card with three-zone layout as specified above")
add_bullet("Hero image loads first — warm, mood-matched to visualMood parameter")
add_bullet("Headline fades in — generated on-device by preference model")
add_bullet("Countdown timer starts — real ticking, mono font")
add_bullet("CTA button pulses once on appear — then static, not aggressive")
add_bullet("Swipe left to dismiss — matches natural gesture")

add_heading("QR Redemption Screen", 3)
add_bullet("Full-screen QR code after Claim is tapped")
add_bullet("Countdown timer visible — same expiry as offer")
add_bullet("On expiry: This moment has passed state replaces QR")
add_bullet("Merchant name and offer summary shown below QR")

add_heading("My Data Screen — Privacy Controls", 3)
add_bullet("Shows all localStorage data in readable plain-English format")
add_bullet("Delete my preference model — one tap, clears localStorage, resets to defaults")
add_bullet("Change location permission tier — links to permission controls")
add_bullet("Data dividend earnings — shows total credits earned from redemptions")
add_bullet("Offer history this month — how many received and redeemed")
add_bullet("Privacy architecture visual: simple device boundary diagram")

add_heading("Merchant Dashboard", 3)
add_body("Three tabs. All data updates in real time via Socket.io.")
add_feature_table(
    ["Tab", "Content"],
    [
        ("Now", "Live transaction velocity vs baseline chart. Active offers count. QUIET PERIOD DETECTED indicator when trigger fires. Current offers on user phones."),
        ("Today", "Offers funnel: generated → accepted → dismissed → expired → redeemed. Revenue from redemptions. Fees charged. Quiet period fill percentage."),
        ("This Week", "Insight card: Your quietest slot: Tuesday 11am–1pm. MOMENTO filled 34% of it this week. Average discount used: 14%. Customers you would not have had: 12. Revenue generated: €54. Fee charged: €2.16."),
    ],
    col_widths=[1.0, 5.5]
)

add_heading("Merchant Setup Screen", 3)
add_body("Three fields. Set once. Never touch again. This is the merchant's entire interaction with MOMENTO.")
add_bullet("What do you offer: dropdown selector with categories")
add_bullet("Maximum discount: slider from 0% to 50%")
add_bullet("Quiet hours to fill: day selector + time range")
add_bullet("Max redemptions per day: number input")
add_bullet("Save and Activate button")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PUSH NOTIFICATION SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 5 — PUSH NOTIFICATION SYSTEM", "Four notification types, timing logic, and the on-device copy architecture")

add_heading("The Core Rule", 2)
add_callout(
    "Maximum 3 notifications per day per user. Send ONLY when all three clocks align. "
    "Never between 10pm and 7am. Never when user is moving fast. Never when an active "
    "offer is already pending. One notification at a time. The offer finds the user — it does not chase them."
)

add_heading("The Permission Journey", 2)
add_body("NEVER show the OS permission dialog on first app open. Show value first.")
add_feature_table(
    ["Step", "What Happens"],
    [
        ("Onboarding complete", "User has already set preferences and seen the app"),
        ("Preview screen", "Show a static mockup of what a real MOMENTO notification looks like"),
        ("Explanation text", "Maximum 3 per day. Only when the moment is genuinely right."),
        ("User choice", "YES, SEND ME MOMENTS button — then OS dialog appears"),
        ("Denial accepted", "If user says no, system accepts completely. In-app experience still works."),
    ],
    col_widths=[1.8, 4.7]
)

add_heading("Four Notification Types", 2)

add_heading("Type 1 — The Moment Offer (Core Product)", 3)
add_body("Trigger: Triple clock alignment detected. Maximum 1 per merchant per user per day. Minimum 2-hour gap between any two notifications.")
add_feature_table(
    ["Element", "Content", "Rule"],
    [
        ("App badge", "Orange dot + MOMENTO", "Brand recognition in 0.3 seconds"),
        ("Timestamp", "now", "Not a time — the word now. Moment-aware language."),
        ("Line 1", "Cold outside? Your cappuccino is waiting.", "Generated on-device by Service Worker — not by server"),
        ("Line 2", "Café Müller · 3 min walk · 11 min left", "Merchant name · walking time · countdown · all in one line"),
        ("Action 1", "CLAIM — 15% OFF", "Discount shown on button — one tap commits"),
        ("Action 2", "Not now", "Dismiss without opening app"),
    ],
    col_widths=[1.5, 2.8, 2.2]
)

add_heading("Type 2 — Expiry Warning", 3)
add_body("Trigger: User tapped CLAIM (QR generated) but has NOT redeemed. Fires exactly once, 3 minutes before QR expiry.")
add_body("Content: Your moment expires in 3 minutes — Café Müller QR code — show on arrival")
add_body("NOT sent if: user already redeemed, user dismissed original offer, timer cancelled by redemption confirmation")

add_heading("Type 3 — Redemption Confirmation", 3)
add_body("Trigger: Payone webhook confirms transaction completed. Fires within 5 seconds of confirmation.")
add_body("Content: Checkmark — Moment redeemed — €0.68 data dividend credit credited — Café Müller Tuesday 11:51am")
add_body("Purpose: Closes the emotional loop. User gets confirmation the system worked. Cashback shown immediately.")

add_heading("Type 4 — City Broadcast", 3)
add_body("Trigger: Merchant quiet period + active Tier 3 and Tier 4 users in city. No precise location known.")
add_body("Content: A quiet moment in Stuttgart right now — Hot drinks · City Centre · 18 min left")
add_body("Differences from Type 1: No distance shown. No merchant name upfront. No tight countdown. SEE THE OFFER instead of CLAIM.")

add_heading("Timing Suppression Logic", 2)
add_body("Before any notification fires, these conditions are checked in order. If any is true, the notification is blocked.")
add_feature_table(
    ["Condition", "Block Reason"],
    [
        ("Time is between 10pm and 7am", "Do Not Disturb hours — hardcoded, no exceptions"),
        ("User movement speed above 15 km/h", "User is driving or on transit — cannot redeem"),
        ("User already has an active unclaimed offer", "One active offer at a time — no stacking ever"),
        ("Last notification sent less than 2 hours ago", "Minimum gap enforcement — no rapid-fire"),
        ("User dismissed last 3 offers from same category", "On-device learning signal — suppress this type"),
        ("Merchant daily redemption cap already reached", "No point sending — merchant cannot honor more"),
    ],
    col_widths=[3.0, 3.5]
)

add_heading("CRITICAL: On-Device Copy Generation", 2)
add_callout(
    "The server NEVER sends the final headline copy. The server sends only parameters. "
    "The Service Worker generates the copy on the device. This is the GDPR compliance mechanism "
    "enforced at the architectural level. Never write code that puts the notification body text "
    "in the server push payload."
)
add_body("What the server sends in the push payload:", bold=True)
add_bullet("emotionalHook: string identifier like cold_weather_warmth")
add_bullet("merchantName: Café Müller")
add_bullet("featuredProduct: Cappuccino + Croissant")
add_bullet("expiryMinutes: 14")
add_bullet("distanceMeters: 280")
add_bullet("offerId: UUID string")
add_bullet("discount: 15")
add_body("What the Service Worker does on the device:", bold=True)
add_bullet("Reads the emotionalHook parameter")
add_bullet("Selects the matching headline template from a local template map")
add_bullet("Fills in the product and merchant values")
add_bullet("Generates: Cold outside? Your Cappuccino + Croissant is waiting.")
add_bullet("Uses that generated string as the notification body")
add_bullet("The generated copy never leaves the device")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PRIVACY ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 6 — PRIVACY ARCHITECTURE", "GDPR compliance enforced in code, not policy documents")

add_heading("The Non-Negotiable Rules", 2)
add_feature_table(
    ["Rule", "Enforcement"],
    [
        ("Raw sensor data never leaves the device", "No API call sends accelerometer, GPS, or motion data"),
        ("Personal preference model never syncs to server", "localStorage weights: no API call ever touches them"),
        ("Claude API receives zero personal identifiers", "Verified in every API call payload before sending"),
        ("When user denies location — no backdoor inference", "No IP lookup, no SSID reading, nothing — code enforced"),
        ("Location shared as city district only", "500m zone string — never precise GPS coordinates"),
        ("Session tokens rotate every 24 hours", "Cannot link behavior across days even with token"),
    ],
    col_widths=[3.0, 3.5]
)

add_heading("What Leaves the Device vs What Stays", 2)
add_feature_table(
    ["Stays on Device FOREVER", "Can Leave Device (Abstract Only)"],
    [
        ("Raw GPS coordinates", "City district string — 500m zone approximation"),
        ("Accelerometer readings", "Intent state: one of three enum values"),
        ("Full browsing history", "Mobility mode: one of three enum values"),
        ("Category preference weights", "Estimated free minutes: rounded number"),
        ("Tone preference weights", "Weather context: enum string like cold_overcast"),
        ("Time slot preference weights", "Anonymous session token: rotates every 24 hours"),
        ("Redeemed offer history", "Nothing else — zero additional fields permitted"),
    ],
    col_widths=[3.0, 3.5]
)

add_heading("Session Token Architecture", 2)
add_body(
    "Every API call uses an anonymous session token generated on the device. The token is a UUID "
    "created in the browser. After 24 hours it is regenerated. The previous token is discarded. "
    "Two days of behavior cannot be linked even if an attacker captured both tokens. "
    "The token identifies a session, not a person."
)

add_heading("The Judge Answer on Privacy", 2)
add_callout(
    '"When a user says no to location — we mean no. There is no code path that infers location after denial. '
    'Our privacy architecture is enforced by the system, not stated in a policy document. '
    'The server never knows what headline the user saw. The server never knows what preference '
    'weights the user has built up. The Claude API receives zero personal identifiers. '
    'GDPR compliance is an architectural property of this system, not a compliance checkbox."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — BUSINESS MODEL & COMPETITIVE
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 7 — BUSINESS MODEL AND COMPETITIVE POSITION", "Revenue model, pricing tiers, and why no competitor can replicate this")

add_heading("Revenue Model — Zero Upfront, Performance Only", 2)
add_feature_table(
    ["Event", "Merchant Pays"],
    [
        ("Offer generated", "€0"),
        ("Offer displayed to user", "€0"),
        ("Offer dismissed by user", "€0"),
        ("Offer expires without redemption", "€0"),
        ("Customer redeems offer at merchant", "3-5% of transaction value only"),
    ],
    col_widths=[3.5, 3.0]
)
add_callout("This is the same model as Uber Eats. The restaurant pays only when an order happens. Risk sits entirely with the platform. The merchant has zero financial risk — if MOMENTO delivers no customers, the merchant pays nothing.")

add_heading("Discount Calculation — Yield Management", 2)
add_body("The AI calculates the MINIMUM effective discount — not the maximum allowed by the merchant.")
add_feature_table(
    ["Scenario", "Discount Logic"],
    [
        ("User is highly receptive + weather is bad + user is very close", "Lower discount — user likely to convert anyway. Merchant saves margin."),
        ("User is moderately browsing + weather is mild + user is further away", "Higher discount — needs more incentive to change course."),
        ("Merchant is extremely quiet (70% below baseline)", "Discount can reach merchant maximum — urgent need justifies it."),
        ("Merchant is mildly slow (35% below baseline)", "Conservative discount — mild nudge only."),
    ],
    col_widths=[3.0, 3.5]
)

add_heading("Three Merchant Tiers", 2)
add_feature_table(
    ["Tier", "Entry Requirement", "Cost Model", "Features"],
    [
        ("Tier 1 — Passive", "Already a Payone customer", "Success fee on redemptions only", "Auto-enrolled, default parameters, basic dashboard"),
        ("Tier 2 — Active", "Logs in once to set rules", "Success fee on redemptions only", "Own rules, richer analytics, pause and resume anytime"),
        ("Tier 3 — Premium", "Monthly subscription fee", "Monthly fee plus lower success rate", "Priority placement, custom branding, advanced forecasting"),
    ],
    col_widths=[1.5, 1.8, 1.8, 1.4]
)

add_heading("Payone — The Unfair Advantage", 2)
add_body("Payone is DSV's payment processing company. Every card payment at a Payone merchant emits a real-time event. MOMENTO is the only system in the world that reads these events to detect commercial quiet periods.")
add_feature_table(
    ["What Payone Provides", "Why No Competitor Has It"],
    [
        ("Real-time transaction velocity per merchant", "Payone is proprietary to DSV — third parties cannot access this data"),
        ("Historical baseline per merchant per time slot", "Requires years of transaction history — cannot be purchased"),
        ("Redemption verification via existing POS", "QR scan at existing terminal = automatic confirmation, no new hardware"),
        ("Automatic success fee collection", "Deducted at point of transaction — no separate billing"),
    ],
    col_widths=[3.0, 3.5]
)
add_body("What Payone cannot do (honest limitations for judges):", bold=True)
add_bullet("Sees transaction events but not what was purchased — product-level needs POS integration")
add_bullet("Merchant-level data only — no individual customer identity within transactions")
add_bullet("Coverage drops outside Germany — primarily a German market asset")
add_bullet("Not every merchant in every street uses Payone — geographic blind spots exist")

add_heading("Sparkassen — The Trust Infrastructure", 2)
add_body("370 independent regional savings banks. 50 million customers. 200+ year embedded relationships in every German town. Legally mandated to serve regional public interest, not maximize profit.")
add_feature_table(
    ["What Sparkassen Provide", "Commercial Impact for MOMENTO"],
    [
        ("Pre-built merchant trust", "Merchant already banks with Sparkasse — MOMENTO inherits 20-year relationships"),
        ("Zero-cost merchant distribution", "Relationship managers already visit merchants — MOMENTO = one added conversation"),
        ("50 million existing users", "Feature inside Sparkasse app = no cold-start download problem"),
        ("Geographic coverage", "Every German town simultaneously — not city-by-city expansion"),
        ("Cultural legitimacy", "Germans trust Sparkasse with their money — trust transfers to MOMENTO"),
    ],
    col_widths=[2.5, 4.0]
)

add_heading("Competitive Position", 2)
add_feature_table(
    ["Competitor", "What They Do", "Their Ceiling", "Our Difference"],
    [
        ("Coupon Apps\n(Stocard, PayPack)", "Digitized paper coupons. Static offers, user must browse.", "Cannot generate offers — only retrieve them. Revenue model trapped in static listings.", "Offers generated at runtime. Never existed before the moment."),
        ("Super-Apps\n(Klarna, Revolut)", "Know what you bought last month. Pre-negotiated chain discounts.", "No real-time merchant quiet signal. Past behavior, not present intent.", "Present intent detected in real time. Only MOMENTO knows a café is quiet NOW."),
        ("Discovery Apps\n(Google Maps)", "Answer the question: where should I go? Requires formed intent.", "No merchant demand signal. GDPR non-compliant in Germany. No local trust.", "Operates before intent forms. The question has not been asked yet."),
        ("Surplus Platforms\n(Too Good To Go)", "End-of-day food surplus. Mystery bags. Food category only.", "Rebuilding for any merchant category would require rebuilding their core product.", "Any quiet period. Any category. AI-generated personalized offers, not mystery bags."),
        ("Ad-Tech\n(Foursquare)", "Sell location-based ad impressions to brands.", "Sell attention, not moments. Charge per impression regardless of conversion.", "Sell moments — transactional offers. Charge only on redemption."),
    ],
    col_widths=[1.4, 1.8, 1.9, 1.4]
)

add_heading("The Amazon Answer", 2)
add_callout(
    '"Our real long-term competitor is Amazon. The reason Amazon cannot build what we are building '
    'is not technical — it is relational. DSV Gruppe sits inside the trust infrastructure of '
    'German local commerce. Amazon sits outside it. German merchants actively distrust Amazon\'s '
    'margin extraction model. Sparkassen have 200-year relationships Amazon cannot purchase. '
    'GDPR enforcement in Germany is the strictest in Europe — Amazon\'s data model is structurally '
    'non-compliant. That gap cannot be purchased."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — THE DEMO
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 8 — THE 90-SECOND DEMO", "Scene breakdown, demo controls, and what must work perfectly")

add_heading("The Demo Scene", 2)
add_body("Stuttgart. Tuesday 11:47am. 8 degrees Celsius feels-like, overcast, 75% rain probability.")

add_feature_table(
    ["Time", "Screen Shown", "What Happens"],
    [
        ("0–15s", "Merchant Dashboard", "Café Müller Payone feed drops. Numbers show: Current 1 transaction, Baseline 12. Quiet Period indicator turns red. System log: QuietPeriodSignal emitted."),
        ("15–25s", "Triple Clock Panel", "Clock 1 turns green: Merchant Quiet. Clock 3 turns green: City Ambient cold and rainy. Clock 2 turns green: User Browsing. All three green: MOMENTO FIRES pulse animation."),
        ("25–45s", "Offer Generating", "Claude API call fires with real parameters. Show parameters streaming in one by one. On-device model writes headline. React component builds widget live: image loads, headline fades in, timer starts."),
        ("45–55s", "User Offer Card", "Full card: Cold outside? Your cappuccino is waiting. Café Müller · 80m · 11 minutes. User taps Warm me up."),
        ("55–65s", "QR Redemption", "QR code unfolds with spring animation. Merchant scans. Dashboard updates in real time: +1 redeemed, €0.18 fee logged, green flash."),
        ("65–80s", "Merchant Dashboard Week View", "Filled 34% of your Tuesday quiet slot this week. 12 customers you would not have had. €54 revenue. €2.16 fee."),
        ("80–90s", "Privacy Architecture", "Simple device vs cloud diagram. When a user says no — we mean no. Tagline: The moment finds the user."),
    ],
    col_widths=[0.8, 1.8, 4.0]
)

add_heading("Demo Controls Panel — Mandatory", 2)
add_body("Visible ONLY when URL contains ?demo=true. Never visible in normal user experience.")
add_feature_table(
    ["Button", "What It Does"],
    [
        ("Trigger Quiet Period", "Forces Payone simulation to drop below quiet threshold immediately"),
        ("Set User to Browsing State", "Forces intent clock green regardless of other signals"),
        ("Fire All Three Clocks", "Triggers all three simultaneously, starts offer generation pipeline"),
        ("Simulate Redemption", "Simulates successful QR scan, updates merchant dashboard via Socket.io"),
        ("Trigger Notification Type 1", "Sends test Moment Offer push notification bypassing timing suppression"),
        ("Trigger Notification Type 3", "Sends test Redemption Confirmation push notification immediately"),
        ("Reset Demo", "Resets all state to baseline, clears offer, resets all clock indicators"),
        ("Speed Control", "1x / 2x / 3x — controls how fast Payone simulation drops to quiet"),
    ],
    col_widths=[2.2, 4.3]
)

add_heading("What Must Work Perfectly — No Exceptions", 2)
add_feature_table(
    ["Feature", "Why It Cannot Fail"],
    [
        ("Triple Clock panel — three indicators turning green live", "This is the hero of the demo. If it does not animate, the core concept is invisible."),
        ("Claude API generates real offer parameters", "Must be a real API call with claude-sonnet-4-6. Cached or mocked responses fail the generative brief requirement."),
        ("Offer widget builds itself live on screen", "Parameters streaming in and widget constructing itself is what makes it generative UI. Static mockup fails."),
        ("Accept flow: QR unfolds from card", "This is the money shot. Spring animation must be smooth and satisfying."),
        ("Merchant dashboard updates via Socket.io on redemption", "Real-time update is the proof of end-to-end loop. Static dashboard fails."),
        ("PWA installable from browser", "Judges must be able to tap Add to Home Screen. If this fails the product is not a PWA."),
    ],
    col_widths=[2.5, 4.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — WHAT NOT TO BUILD
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 9 — WHAT NOT TO BUILD", "Save time by knowing exactly what to skip for the hackathon")

add_heading("Explicitly Out of Scope for Hackathon", 2)
add_callout("The brief says: This challenge is won in the interaction design, not the model architecture. Do not over-engineer the AI stack and under-engineer the experience.")

add_feature_table(
    ["Do NOT Build", "Use Instead", "Why"],
    [
        ("Real TensorFlow Lite on-device model", "React state + localStorage + demo toggle", "Weeks of work. Zero visible difference in demo."),
        ("Real Payone API integration", "Node.js setInterval with realistic transaction patterns", "No API access exists. Simulate convincingly."),
        ("Real GPS tracking and accelerometer", "Location selector or hardcoded Stuttgart + demo toggle", "Permission flows take time. Simulation is cleaner for demo."),
        ("Authentication and user accounts", "Anonymous session tokens only", "No judge needs to log in to evaluate the product."),
        ("Production database", "In-memory Maps and arrays in Node.js", "Zero setup time. Zero configuration. Works perfectly for demo."),
        ("Real on-device SLM like Phi-3", "localStorage template map with preference weights", "Model download alone would fail the demo. Simulate the output."),
        ("App Store submission", "PWA at Vercel URL", "No review time. Instant access for judges."),
        ("Multi-city support", "Hardcode Stuttgart with configurable thresholds", "Context.config.json shows it is configurable. Demo one city perfectly."),
    ],
    col_widths=[2.0, 2.0, 2.5]
)

add_heading("Build Priority Order", 2)
add_feature_table(
    ["Priority", "Feature", "Why This Order"],
    [
        ("P0 — Demo breaks without this", "Triple Clock panel with live Payone simulation", "Visual centerpiece. Everything else demonstrates from here."),
        ("P0 — Demo breaks without this", "Claude API offer generation with claude-sonnet-4-6", "Test this FIRST before any UI. Confirm JSON response works."),
        ("P0 — Demo breaks without this", "Offer card with GenUI widget building live", "The product moment. Hero image, headline, countdown, CTA."),
        ("P0 — Demo breaks without this", "Accept → QR code with spring animation", "The satisfying completion moment. Must feel premium."),
        ("P0 — Demo breaks without this", "Merchant dashboard with Socket.io real-time update", "Proves end-to-end loop. Counter increments on redemption."),
        ("P0 — Demo breaks without this", "PWA manifest and installable from browser", "Required deliverable in brief."),
        ("P1 — Strong submission", "Dismiss animation with branded copy", "Judges will test dismiss. Must say 'Another moment is coming.'"),
        ("P1 — Strong submission", "Expire animation with branded copy", "Judges will let timer run. Must say 'This moment has passed.'"),
        ("P1 — Strong submission", "Onboarding 5-step flow", "Shows privacy-first design thinking."),
        ("P1 — Strong submission", "Push notification system — all 4 types", "Required PWA feature. Service Worker on-device copy generation."),
        ("P1 — Strong submission", "Privacy architecture diagram in app", "Judge question about GDPR must have visual answer."),
        ("P2 — Nice to have", "Real OpenWeatherMap weather widget", "Shows real context data. Adds credibility."),
        ("P2 — Nice to have", "My Data screen with delete option", "Shows genuine privacy commitment."),
        ("P2 — Nice to have", "Multiple demo merchants", "Richens the demo but one merchant done perfectly is sufficient."),
    ],
    col_widths=[2.2, 2.5, 1.8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — JUDGE Q&A RAPID FIRE
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 10 — ANTICIPATED JUDGE QUESTIONS", "Every likely question and the exact answer to give")

add_heading("Technical Questions", 2)
add_feature_table(
    ["Judge Asks", "Answer"],
    [
        ("Show me the code where notification copy is generated on-device", "Service Worker push event listener contains generateHeadlineFromParams function. Server sends emotionalHook parameter. Device generates the string. Server never knows what copy was shown."),
        ("Is this TensorFlow Lite or a toggle button?", "Demo uses a toggle for clarity — judges can see exact causation. Production path is TF Lite on-device. We show the mechanism visually, not hidden inside sensor data."),
        ("What Claude model are you using?", "claude-sonnet-4-6. Verify this is correct in every API call before demo."),
        ("What happens to the user if they walk away after claiming?", "QR code remains valid for its full expiry duration regardless of movement. User committed — we do not punish them for moving."),
        ("How do you prevent QR code fraud?", "HMAC-SHA256 signature on every token. Single-use enforcement via used-token Set in memory. Expiry enforced server-side. Screenshot sharing triggers already-used error."),
        ("Why not build a native app?", "PWA achieves everything in the brief. Judges access via URL in 5 seconds. No App Store review delay. We spent development time on the product, not the framework."),
    ],
    col_widths=[2.5, 4.0]
)

add_heading("Privacy and GDPR Questions", 2)
add_feature_table(
    ["Judge Asks", "Answer"],
    [
        ("What if the user denies location permission?", "System accepts that completely. No IP lookup. No SSID reading. No backdoor. Tier 3: uses only the neighborhood text user typed manually. Tier 4: no location features at all."),
        ("How is this GDPR compliant?", "Raw sensor data never leaves device. Preference model in localStorage only, never synced. Claude API receives zero personal identifiers. Session tokens rotate every 24 hours. Compliance is architectural, not a policy document."),
        ("What personal data do you store?", "Server stores nothing personal. localStorage stores preference weights that the user can delete in one tap. Session token rotates daily. No user account, no profile, no persistent identity."),
    ],
    col_widths=[2.5, 4.0]
)

add_heading("Business Questions", 2)
add_feature_table(
    ["Judge Asks", "Answer"],
    [
        ("Do merchants pay upfront?", "Zero upfront. Zero monthly for Tier 1 and 2. Success fee of 3-5% only when a customer actually redeems. Merchant pays nothing if MOMENTO delivers nothing."),
        ("Do users need to download an app?", "No App Store download. PWA accessed via URL. Long term: feature inside Sparkasse app already on 50 million German phones."),
        ("What if no merchants are in Payone?", "Acknowledged limitation. Payone covers significant portion of German local merchants. Coverage gaps are honest — system only detects quiet periods for Payone merchants."),
        ("Who are your competitors?", "Five categories: coupon apps, super-apps, discovery apps, surplus platforms, ad-tech. None have the combination of real-time merchant demand signal from Payone plus Sparkassen trust plus generative offers. Too Good To Go is philosophically closest but food surplus only."),
        ("Why can't Amazon build this?", "Not a technical limitation. A relational one. German merchants distrust Amazon's margin extraction model. Sparkassen have 200-year embedded relationships Amazon cannot purchase. GDPR in Germany blocks Amazon's data model structurally."),
        ("How does the discount amount get decided?", "Minimum effective discount via yield management. More receptive user = lower discount needed. More urgent merchant quiet period = higher discount justified. Same logic airlines use for dynamic pricing."),
    ],
    col_widths=[2.5, 4.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — ENVIRONMENT & DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 11 — ENVIRONMENT VARIABLES AND DEPLOYMENT", "Everything needed to run the project")

add_heading("Required Environment Variables", 2)
add_feature_table(
    ["Variable", "Source", "Used For"],
    [
        ("ANTHROPIC_API_KEY", "Anthropic console — console.anthropic.com", "Claude API calls for offer generation"),
        ("OPENWEATHERMAP_API_KEY", "openweathermap.org free tier", "Real weather data for Stuttgart context"),
        ("TOKEN_SECRET", "Generate: openssl rand -hex 32", "HMAC-SHA256 signing of QR redemption tokens"),
        ("VAPID_PUBLIC_KEY", "Generate: npx web-push generate-vapid-keys", "Web Push API public key for browsers"),
        ("VAPID_PRIVATE_KEY", "Same generation command as above", "Web Push API private key for server signing"),
        ("VAPID_SUBJECT", "Your email or project URL", "Push notification sender identification"),
        ("PORT", "3001 or any available port", "Express server port"),
    ],
    col_widths=[2.0, 2.2, 2.3]
)

add_heading("Deployment Steps", 2)
add_bullet("Install Vercel CLI: npm install -g vercel")
add_bullet("From project root: vercel deploy")
add_bullet("Add environment variables in Vercel dashboard under Settings → Environment Variables")
add_bullet("Vercel provides automatic HTTPS — required for PWA push notifications and geolocation API")
add_bullet("Custom domain: register momento.app and point to Vercel deployment")

add_heading("First Thing to Test Before Building Anything", 2)
add_callout(
    "BEFORE WRITING ANY UI CODE: Make one test Claude API call using the model claude-sonnet-4-6 "
    "with a simple JSON output prompt. Confirm you get a valid JSON response. If this fails, "
    "nothing else in the project will work. Fix the API key and model string first. Only then "
    "start building components."
)

add_heading("Demo URL Structure", 2)
add_bullet("Normal user experience: momento.app/")
add_bullet("Demo controls panel visible: momento.app/?demo=true")
add_bullet("Merchant dashboard: momento.app/merchant")
add_bullet("Merchant setup: momento.app/merchant/setup")
add_bullet("Privacy controls: momento.app/my-data")
add_bullet("Onboarding: momento.app/onboarding")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — FINAL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_section_banner("SECTION 12 — THE ONE-PARAGRAPH PITCH", "What to say when you have 30 seconds")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(16)
r = p.add_run(
    "Every day, thousands of local merchants sit through quiet hours while receptive customers walk past, "
    "completely unaware. MOMENTO is the layer that connects them — not by showing users a list of deals "
    "to browse, but by detecting the exact moment a merchant needs a customer, finding a receptive person "
    "nearby, and generating a single, specific, time-sensitive offer before the user even knew they wanted it. "
    "We built this on DSV's existing infrastructure: Payone tells us which merchant is quiet right now — "
    "a signal no other system in the world has. Sparkassen get us into that merchant's door through trust "
    "relationships built over 200 years. And they put MOMENTO in front of 50 million Germans through an app "
    "they already trust with their money. No startup can buy these assets. No tech giant can culturally earn them. "
    "MOMENTO does not just give local merchants better marketing. It gives them the same anticipatory intelligence "
    "Amazon built for global commerce — pointed at the corner café, respecting privacy by architecture, "
    "and charging only when it actually works."
)
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY

add_divider()

add_heading("The Tagline", 2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The moment finds the user.")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = AMBER

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("MOMENTO — The Generative City-Wallet")
r2.font.size = Pt(13)
r2.font.color.rgb = NAVY

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/mnt/user-data/outputs/MOMENTO_Complete_Project_Overview.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
